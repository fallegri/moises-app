"""Tests for file parsing service."""

import io
import pytest

from docx import Document as DocxDocument
from openpyxl import Workbook

from app.services.file_parser import FileParserService


@pytest.fixture
def parser():
    """Create a file parser service instance."""
    return FileParserService()


@pytest.fixture
def sample_docx() -> bytes:
    """Create a sample .docx file as bytes."""
    doc = DocxDocument()
    doc.add_paragraph("This is the first paragraph.")
    doc.add_paragraph("This is the second paragraph.")
    doc.add_paragraph("")  # empty paragraph
    doc.add_paragraph("This is the third paragraph with more content.")
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


@pytest.fixture
def sample_xlsx() -> bytes:
    """Create a sample .xlsx file as bytes."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Data"
    ws.append(["Name", "Age", "City"])
    ws.append(["Alice", 30, "Lima"])
    ws.append(["Bob", 25, "Cusco"])
    ws.append([None, None, None])  # empty row
    ws.append(["Charlie", 35, "Arequipa"])
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


@pytest.fixture
def sample_md() -> bytes:
    """Create a sample markdown file as bytes."""
    content = """# Research Title

## Introduction

This is the introduction section.

## Methodology

This describes the methodology used.
"""
    return content.encode("utf-8")


class TestFileParser:
    """Tests for FileParserService."""

    def test_parse_docx(self, parser, sample_docx):
        """Test parsing a .docx file."""
        result = parser.parse_file(sample_docx, "document.docx")
        assert "first paragraph" in result
        assert "second paragraph" in result
        assert "third paragraph" in result

    def test_parse_xlsx(self, parser, sample_xlsx):
        """Test parsing a .xlsx file."""
        result = parser.parse_file(sample_xlsx, "data.xlsx")
        assert "Alice" in result
        assert "Bob" in result
        assert "Lima" in result
        assert "Cusco" in result
        assert "Hoja: Data" in result

    def test_parse_markdown(self, parser, sample_md):
        """Test parsing a markdown file."""
        result = parser.parse_file(sample_md, "document.md")
        assert "Research Title" in result
        assert "Introduction" in result
        assert "Methodology" in result

    def test_parse_txt(self, parser):
        """Test parsing a plain text file."""
        content = b"This is plain text content.\nWith multiple lines."
        result = parser.parse_file(content, "notes.txt")
        assert "plain text content" in result
        assert "multiple lines" in result

    def test_unsupported_format(self, parser):
        """Test that unsupported format raises ValueError."""
        with pytest.raises(ValueError, match="Unsupported file format"):
            parser.parse_file(b"content", "file.pdf")

    def test_supported_extensions(self, parser):
        """Test supported extensions set."""
        assert ".docx" in parser.SUPPORTED_EXTENSIONS
        assert ".xlsx" in parser.SUPPORTED_EXTENSIONS
        assert ".md" in parser.SUPPORTED_EXTENSIONS
        assert ".txt" in parser.SUPPORTED_EXTENSIONS
        assert ".pdf" not in parser.SUPPORTED_EXTENSIONS

    def test_docx_empty_paragraphs_skipped(self, parser):
        """Test that empty paragraphs in docx are skipped."""
        doc = DocxDocument()
        doc.add_paragraph("")
        doc.add_paragraph("Content")
        doc.add_paragraph("")
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        result = parser.parse_file(buffer.getvalue(), "test.docx")
        assert result.strip() == "Content"

    def test_xlsx_multiple_sheets(self, parser):
        """Test parsing xlsx with multiple sheets."""
        wb = Workbook()
        ws1 = wb.active
        ws1.title = "Sheet1"
        ws1.append(["Data1"])

        ws2 = wb.create_sheet("Sheet2")
        ws2.append(["Data2"])

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)

        result = parser.parse_file(buffer.getvalue(), "multi.xlsx")
        assert "Sheet1" in result
        assert "Sheet2" in result
        assert "Data1" in result
        assert "Data2" in result

    def test_markdown_utf8_encoding(self, parser):
        """Test that markdown files handle UTF-8 characters."""
        content = "Investigacion en espanol: acentos y enes".encode("utf-8")
        result = parser.parse_file(content, "spanish.md")
        assert "espanol" in result
        assert "acentos" in result
