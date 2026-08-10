"""Document generator service for APA 7 formatted Word documents."""

import io
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


class DocumentGeneratorService:
    """Generates APA 7 formatted Word documents."""

    def __init__(self):
        self.font_name = "Times New Roman"
        self.font_size = Pt(12)
        self.line_spacing = 2.0  # Double spacing
        self.margin_top = Inches(1)
        self.margin_bottom = Inches(1)
        self.margin_left = Inches(1)
        self.margin_right = Inches(1)

    def create_document(
        self,
        title: str,
        content: str,
        author: Optional[str] = None,
        references: Optional[list[str]] = None,
    ) -> bytes:
        """Create an APA 7 formatted Word document.

        Args:
            title: Document title.
            content: Main body content.
            author: Optional author name.
            references: Optional list of references.

        Returns:
            Document as bytes.
        """
        doc = Document()
        self._set_margins(doc)
        self._setup_styles(doc)

        # Title page
        self._add_title_page(doc, title, author)

        # Main content
        self._add_content(doc, content)

        # References
        if references:
            self._add_references(doc, references)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _set_margins(self, doc: Document):
        """Set APA 7 margins (1 inch on all sides)."""
        sections = doc.sections
        for section in sections:
            section.top_margin = self.margin_top
            section.bottom_margin = self.margin_bottom
            section.left_margin = self.margin_left
            section.right_margin = self.margin_right

    def _setup_styles(self, doc: Document):
        """Configure document styles for APA 7."""
        style = doc.styles["Normal"]
        font = style.font
        font.name = self.font_name
        font.size = self.font_size
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = self.line_spacing
        paragraph_format.space_after = Pt(0)
        paragraph_format.space_before = Pt(0)

    def _add_title_page(self, doc: Document, title: str, author: Optional[str] = None):
        """Add APA 7 title page."""
        # Add blank lines for centering
        for _ in range(4):
            doc.add_paragraph("")

        # Title - centered, bold
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title)
        run.bold = True
        run.font.name = self.font_name
        run.font.size = self.font_size

        # Author
        if author:
            doc.add_paragraph("")
            author_para = doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = author_para.add_run(author)
            run.font.name = self.font_name
            run.font.size = self.font_size

        # Page break after title page
        doc.add_page_break()

    def _add_content(self, doc: Document, content: str):
        """Add main body content with proper formatting."""
        lines = content.split("\n")

        for line in lines:
            line = line.strip()
            if not line:
                continue

            if line.startswith("### "):
                # Level 3 heading - flush left, bold italic
                para = doc.add_paragraph()
                run = para.add_run(line[4:])
                run.bold = True
                run.italic = True
                run.font.name = self.font_name
                run.font.size = self.font_size
            elif line.startswith("## "):
                # Level 2 heading - flush left, bold
                para = doc.add_paragraph()
                run = para.add_run(line[3:])
                run.bold = True
                run.font.name = self.font_name
                run.font.size = self.font_size
            elif line.startswith("# "):
                # Level 1 heading - centered, bold
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(line[2:])
                run.bold = True
                run.font.name = self.font_name
                run.font.size = self.font_size
            else:
                # Normal paragraph with first-line indent
                para = doc.add_paragraph()
                para.paragraph_format.first_line_indent = Inches(0.5)
                run = para.add_run(line)
                run.font.name = self.font_name
                run.font.size = self.font_size

    def _add_references(self, doc: Document, references: list[str]):
        """Add references section following APA 7 format."""
        doc.add_page_break()

        # References heading - centered, bold
        heading_para = doc.add_paragraph()
        heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading_para.add_run("Referencias")
        run.bold = True
        run.font.name = self.font_name
        run.font.size = self.font_size

        # Each reference with hanging indent
        for ref in references:
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.first_line_indent = Inches(-0.5)
            run = para.add_run(ref)
            run.font.name = self.font_name
            run.font.size = self.font_size
